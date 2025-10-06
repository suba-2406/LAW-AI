from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.utils.timezone import now
from django.views.decorators.http import require_POST

from .forms import PDFUploadForm, RentalAgreementForm, DivorceAgreementForm, LandAgreementForm
from .models import UploadedPDF, ChatHistory
from .utils import process_pdf_and_summarize, extract_text_from_pdf

from bardapi import Bard
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert

import os
import uuid
import time
import pythoncom

from django.http import HttpResponseServerError, FileResponse
import tempfile
from io import BytesIO

def home(request):
    return render(request, 'core/home.html')


def upload_pdf(request):
    summary = None
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            pdf = form.save()
            summary = process_pdf_and_summarize(pdf.file.path)
            pdf.summary = summary
            pdf.save()
    else:
        form = PDFUploadForm()
    return render(request, 'core/upload.html', {'form': form, 'summary': summary})


def chat_about_pdf(request, pdf_id):
    pdf = get_object_or_404(UploadedPDF, id=pdf_id)
    question = request.GET.get("q")
    answer = ""
    if question:
        full_text = extract_text_from_pdf(pdf.file.path)
        context = " ".join(full_text[:3])
        prompt = f"Based on this document:\n'''{context}'''\nAnswer this: {question}"
        try:
            answer = Bard().get_answer(prompt)['content']
        except:
            answer = "Could not process the question."

        # Save to chat history
        ChatHistory.objects.create(pdf=pdf, question=question, answer=answer)

    chat_history = pdf.chats.order_by('-asked_at')
    return render(request, 'core/chat.html', {
        'pdf': pdf,
        'answer': answer,
        'chat_history': chat_history
    })


def download_summary(request, pdf_id):
    pdf = get_object_or_404(UploadedPDF, id=pdf_id)
    response = HttpResponse(pdf.summary, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename=summary.txt'
    return response


def document_selector(request):
    return render(request, 'law_documents/document_selector.html')


def generate_doc(request, form_class, template_name, doc_name, html_template):
    if request.method == 'POST':
        form = form_class(request.POST)
        if form.is_valid():
            doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', template_name)
            doc = DocxTemplate(doc_template_path)
            doc.render(form.cleaned_data)

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{doc_name}.docx"'
            doc.save(response)
            return response
    else:
        form = form_class()
    return render(request, f'law_documents/{html_template}', {'form': form})


def generate_rental_agreement(request):
    return generate_doc(request, RentalAgreementForm, 'rental_template.docx', 'Rental_Agreement', 'generate_rental.html')


def generate_divorce_agreement(request):
    return generate_doc(request, DivorceAgreementForm, 'divorce_template.docx', 'Divorce_Agreement', 'generate_divorce.html')


def generate_land_agreement(request):
    return generate_doc(request, LandAgreementForm, 'land_template.docx', 'Land_Agreement', 'generate_land.html')


@require_POST
def preview_rental_doc(request):
    form = RentalAgreementForm(request.POST)
    if not form.is_valid():
        # Return errors as JSON or simple message
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'rental_template.docx')
    if not os.path.exists(doc_template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    # Load template and render with form data
    doc = DocxTemplate(doc_template_path)
    doc.render(form.cleaned_data)

    # Save rendered doc to in-memory file
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Load rendered doc to extract text
    document = Document(doc_io)
    content = []

    # Extract paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    # Join paragraphs with <br> for HTML display
    rendered_html = "<br>".join(content)

    # Return JSON with HTML content
    return JsonResponse({'rendered_html': rendered_html})



@require_POST
def preview_divorce_doc(request):
    form = DivorceAgreementForm(request.POST)
    if not form.is_valid():
        print("Form errors:", form.errors)  # Add this for debugging
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    doc_template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'divorce_template.docx')
    if not os.path.exists(doc_template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    doc = DocxTemplate(doc_template_path)
    doc.render(form.cleaned_data)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    document = Document(doc_io)
    content = []

    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    rendered_html = "<br>".join(content)
    return JsonResponse({'rendered_html': rendered_html})

@require_POST
def preview_land_doc(request):
    form = LandAgreementForm(request.POST)
    if not form.is_valid():
        return JsonResponse({'error': 'Invalid data submitted.'}, status=400)

    template_path = os.path.join(settings.BASE_DIR, 'core', 'docx_templates', 'land_template.docx')
    if not os.path.exists(template_path):
        return JsonResponse({'error': 'Template not found.'}, status=404)

    doc = DocxTemplate(template_path)
    doc.render(form.cleaned_data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    document = Document(buffer)
    content = []

    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    return JsonResponse({'rendered_html': "<br>".join(content)})


